from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass
from datetime import timedelta
from pathlib import Path
from typing import Any
from uuid import UUID

from docx import Document
from sqlalchemy import and_, delete, or_, select
from sqlalchemy.orm import Session

from ..adapters.object_storage import get_object_storage
from ..database import SessionLocal, set_public_tenant_context
from ..model_mixins import utcnow
from ..support_ai_models import (
    SupportAIIngestionJobRow,
    SupportAIKnowledgeChunkRow,
    SupportAIKnowledgeSourceRow,
)
from .embedding import EmbeddingProviderError, precompute_embeddings
from .embedding_configuration import resolved_text_embedding_provider
from .support_ai_language import detect_message_language


MAX_EXTRACTED_CHARACTERS = 2_500_000
MAX_CHUNKS_PER_SOURCE = 2500
TARGET_CHUNK_CHARACTERS = 1400
CHUNK_OVERLAP_CHARACTERS = 180


@dataclass(frozen=True, slots=True)
class ParsedKnowledgeBlock:
    text: str
    section_path: str
    locator: dict[str, Any]


class KnowledgeIngestionError(ValueError):
    def __init__(self, code: str, message: str) -> None:
        super().__init__(message)
        self.code = code


def claim_next_knowledge_ingestion(
    session: Session,
    *,
    tenant_id: UUID,
    stale_after_seconds: int = 900,
) -> tuple[UUID, UUID] | None:
    """Claim one durable ingestion job, including work abandoned by a crash."""

    now = utcnow()
    stale_before = now - timedelta(seconds=max(60, stale_after_seconds))
    job = session.scalar(
        select(SupportAIIngestionJobRow)
        .where(
            SupportAIIngestionJobRow.tenant_id == tenant_id,
            or_(
                SupportAIIngestionJobRow.status == "QUEUED",
                and_(
                    SupportAIIngestionJobRow.status == "RUNNING",
                    or_(
                        SupportAIIngestionJobRow.started_at.is_(None),
                        SupportAIIngestionJobRow.started_at <= stale_before,
                    ),
                ),
            ),
        )
        .order_by(SupportAIIngestionJobRow.created_at, SupportAIIngestionJobRow.id)
        .with_for_update(skip_locked=True)
        .limit(1)
    )
    if job is None:
        session.rollback()
        return None
    job.status = "RUNNING"
    job.started_at = now
    session.commit()
    return job.source_id, job.id


def _clean_text(value: str) -> str:
    value = value.replace("\x00", " ")
    value = re.sub(r"[\t\f\v]+", " ", value)
    value = re.sub(r"[ ]{2,}", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def _decode_text(path: Path) -> str:
    payload = path.read_bytes()
    if len(payload) > MAX_EXTRACTED_CHARACTERS * 4:
        raise KnowledgeIngestionError("KNOWLEDGE_FILE_TOO_LARGE", "文本文件内容过大。")
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise KnowledgeIngestionError(
        "KNOWLEDGE_TEXT_ENCODING_UNSUPPORTED",
        "文本文件必须使用 UTF-8 或 GB18030 编码。",
    )


def _parse_text(path: Path) -> tuple[list[ParsedKnowledgeBlock], str, str]:
    raw = _decode_text(path)
    blocks: list[ParsedKnowledgeBlock] = []
    heading = path.stem
    paragraph_index = 0
    for piece in re.split(r"\n\s*\n", raw):
        text = _clean_text(piece)
        if not text:
            continue
        if path.suffix.casefold() == ".md" and re.match(r"^#{1,6}\s+", text):
            heading = re.sub(r"^#{1,6}\s+", "", text).strip() or heading
            continue
        paragraph_index += 1
        blocks.append(
            ParsedKnowledgeBlock(
                text=text,
                section_path=heading,
                locator={"type": "paragraph", "paragraph": paragraph_index},
            )
        )
    return blocks, "plain-text", "1"


def _parse_docx(path: Path) -> tuple[list[ParsedKnowledgeBlock], str, str]:
    document = Document(path)
    blocks: list[ParsedKnowledgeBlock] = []
    heading = path.stem
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = _clean_text(paragraph.text)
        if not text:
            continue
        style_name = str(paragraph.style.name or "").casefold()
        if style_name.startswith("heading") or style_name.startswith("标题"):
            heading = text
            continue
        blocks.append(
            ParsedKnowledgeBlock(
                text=text,
                section_path=heading,
                locator={"type": "docx_paragraph", "paragraph": index},
            )
        )
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            text = _clean_text(" | ".join(cell.text for cell in row.cells))
            if text:
                blocks.append(
                    ParsedKnowledgeBlock(
                        text=text,
                        section_path=f"{heading} / 表格 {table_index}",
                        locator={
                            "type": "docx_table",
                            "table": table_index,
                            "row": row_index,
                        },
                    )
                )
    return blocks, "python-docx", "1"


def _parse_pdf(path: Path) -> tuple[list[ParsedKnowledgeBlock], str, str]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - deployment dependency guard
        raise KnowledgeIngestionError(
            "KNOWLEDGE_PDF_PARSER_UNAVAILABLE",
            "PDF 解析组件尚未安装，请联系平台管理员。",
        ) from exc
    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        raise KnowledgeIngestionError(
            "KNOWLEDGE_PDF_INVALID", "PDF 文件损坏或已加密，无法解析。"
        ) from exc
    if reader.is_encrypted:
        try:
            if reader.decrypt("") == 0:
                raise KnowledgeIngestionError(
                    "KNOWLEDGE_PDF_ENCRYPTED", "暂不支持需要密码的 PDF。"
                )
        except KnowledgeIngestionError:
            raise
        except Exception as exc:
            raise KnowledgeIngestionError(
                "KNOWLEDGE_PDF_ENCRYPTED", "暂不支持需要密码的 PDF。"
            ) from exc
    blocks: list[ParsedKnowledgeBlock] = []
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        for paragraph_index, piece in enumerate(
            re.split(r"\n\s*\n|(?<=。)\s*\n", page_text), start=1
        ):
            text = _clean_text(piece)
            if text:
                blocks.append(
                    ParsedKnowledgeBlock(
                        text=text,
                        section_path=f"第 {page_number} 页",
                        locator={
                            "type": "pdf_page",
                            "page": page_number,
                            "paragraph": paragraph_index,
                        },
                    )
                )
    return blocks, "pypdf", "1"


def parse_knowledge_file(
    path: Path,
    *,
    original_filename: str,
) -> tuple[list[ParsedKnowledgeBlock], str, str]:
    suffix = Path(original_filename).suffix.casefold()
    if suffix in {".txt", ".md"}:
        blocks, parser, version = _parse_text(path)
    elif suffix == ".docx":
        blocks, parser, version = _parse_docx(path)
    elif suffix == ".pdf":
        blocks, parser, version = _parse_pdf(path)
    else:
        raise KnowledgeIngestionError(
            "KNOWLEDGE_FILE_TYPE_UNSUPPORTED",
            "仅支持 PDF、DOCX、TXT 和 Markdown 文件。",
        )
    total = sum(len(block.text) for block in blocks)
    if total <= 0:
        raise KnowledgeIngestionError(
            "KNOWLEDGE_FILE_EMPTY",
            "文件中没有可用于知识库的文本；扫描版 PDF 请先执行 OCR。",
        )
    if total > MAX_EXTRACTED_CHARACTERS:
        raise KnowledgeIngestionError(
            "KNOWLEDGE_CONTENT_TOO_LARGE",
            "文件提取后的文本超过当前知识库上限。",
        )
    return blocks, parser, version


def _split_long_text(text: str, maximum: int) -> list[str]:
    if len(text) <= maximum:
        return [text]
    pieces: list[str] = []
    cursor = 0
    while cursor < len(text):
        end = min(len(text), cursor + maximum)
        if end < len(text):
            boundary = max(
                text.rfind("。", cursor, end),
                text.rfind(". ", cursor, end),
                text.rfind("\n", cursor, end),
            )
            if boundary > cursor + maximum // 2:
                end = boundary + 1
        pieces.append(text[cursor:end].strip())
        if end >= len(text):
            break
        cursor = max(end - CHUNK_OVERLAP_CHARACTERS, cursor + 1)
    return [piece for piece in pieces if piece]


def build_knowledge_chunks(
    blocks: list[ParsedKnowledgeBlock],
) -> list[ParsedKnowledgeBlock]:
    chunks: list[ParsedKnowledgeBlock] = []
    current_texts: list[str] = []
    current_section = "正文"
    current_locator: dict[str, Any] = {}

    def flush() -> None:
        nonlocal current_texts, current_section, current_locator
        if not current_texts:
            return
        content = "\n\n".join(current_texts).strip()
        if content:
            chunks.append(
                ParsedKnowledgeBlock(
                    text=content,
                    section_path=current_section,
                    locator=dict(current_locator),
                )
            )
        current_texts = []
        current_locator = {}

    for block in blocks:
        for piece in _split_long_text(block.text, TARGET_CHUNK_CHARACTERS):
            projected_size = sum(len(value) for value in current_texts) + len(piece)
            if current_texts and (
                projected_size > TARGET_CHUNK_CHARACTERS
                or block.section_path != current_section
            ):
                flush()
            if not current_texts:
                current_section = block.section_path
                current_locator = dict(block.locator)
            current_texts.append(piece)
    flush()
    if len(chunks) > MAX_CHUNKS_PER_SOURCE:
        raise KnowledgeIngestionError(
            "KNOWLEDGE_CHUNK_LIMIT_EXCEEDED",
            "文件分块数量过多，请拆分成多个知识文件后重新上传。",
        )
    return chunks


def _sha256(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def index_knowledge_source(
    session: Session,
    *,
    source: SupportAIKnowledgeSourceRow,
    job: SupportAIIngestionJobRow,
) -> None:
    previously_available = (
        source.status in {"READY", "APPROVED"} and source.chunk_count > 0
    )
    job.status = "RUNNING"
    job.progress = 5
    job.started_at = utcnow()
    if not previously_available:
        source.status = "PROCESSING"
    source.failure_code = None
    source.failure_message = None
    session.commit()
    media = source.media_object_id
    from ..file_security_models import MediaObjectRow

    media_row = session.scalar(
        select(MediaObjectRow).where(
            MediaObjectRow.tenant_id == source.tenant_id,
            MediaObjectRow.id == media,
            MediaObjectRow.status == "AVAILABLE",
        )
    )
    if media_row is None:
        raise KnowledgeIngestionError(
            "KNOWLEDGE_MEDIA_UNAVAILABLE", "知识文件不存在或当前不可用。"
        )
    with get_object_storage().materialize(media_row.object_key) as path:
        blocks, parser, parser_version = parse_knowledge_file(
            path,
            original_filename=source.original_filename,
        )
    job.parser_identifier = parser
    job.parser_version = parser_version
    job.progress = 35
    chunks = build_knowledge_chunks(blocks)
    texts = [chunk.text for chunk in chunks]
    try:
        embedder = resolved_text_embedding_provider(session)
        precomputed = precompute_embeddings(embedder, texts, batch_size=64)
        vectors = precomputed.embed(texts)
    except (EmbeddingProviderError, ValueError) as exc:
        raise KnowledgeIngestionError(
            "KNOWLEDGE_EMBEDDING_FAILED",
            "文件解析成功，但向量化失败，请检查配置中心的 Embedding 配置。",
        ) from exc
    job.progress = 70
    session.execute(
        delete(SupportAIKnowledgeChunkRow).where(
            SupportAIKnowledgeChunkRow.tenant_id == source.tenant_id,
            SupportAIKnowledgeChunkRow.source_id == source.id,
        )
    )
    detected_language = source.language
    if detected_language == "und":
        detected_language = detect_message_language("\n".join(texts[:12]))
    for index, (chunk, vector) in enumerate(zip(chunks, vectors, strict=True)):
        session.add(
            SupportAIKnowledgeChunkRow(
                tenant_id=source.tenant_id,
                source_id=source.id,
                chunk_index=index,
                section_path=chunk.section_path[:500],
                content=chunk.text,
                content_hash=_sha256(chunk.text),
                token_count=max(1, len(chunk.text) // 4),
                language=detected_language,
                locator=chunk.locator,
                embedding=[float(value) for value in vector],
                embedding_provider=embedder.identity.provider,
                embedding_model=embedder.identity.model_name,
                embedding_version=embedder.identity.model_version,
                embedding_dimensions=embedder.identity.dimensions,
                status="ACTIVE",
            )
        )
    source.language = detected_language
    source.chunk_count = len(chunks)
    source.version += 1
    source.status = "APPROVED"
    source.approved_at = utcnow()
    source.approved_by_user_id = source.created_by_user_id
    job.status = "SUCCEEDED"
    job.progress = 100
    job.chunks_written = len(chunks)
    job.completed_at = utcnow()
    session.commit()


def process_knowledge_ingestion(
    *,
    tenant_id: UUID,
    source_id: UUID,
    job_id: UUID,
) -> None:
    with SessionLocal() as session:
        set_public_tenant_context(session, tenant_id=tenant_id)
        source = session.scalar(
            select(SupportAIKnowledgeSourceRow).where(
                SupportAIKnowledgeSourceRow.tenant_id == tenant_id,
                SupportAIKnowledgeSourceRow.id == source_id,
            )
        )
        job = session.scalar(
            select(SupportAIIngestionJobRow).where(
                SupportAIIngestionJobRow.tenant_id == tenant_id,
                SupportAIIngestionJobRow.id == job_id,
                SupportAIIngestionJobRow.source_id == source_id,
            )
        )
        if source is None or job is None or job.status not in {"QUEUED", "RUNNING"}:
            return
        try:
            index_knowledge_source(session, source=source, job=job)
        except KnowledgeIngestionError as exc:
            session.rollback()
            source = session.get(SupportAIKnowledgeSourceRow, source_id)
            job = session.get(SupportAIIngestionJobRow, job_id)
            if source is not None:
                if source.approved_at is None or source.chunk_count <= 0:
                    source.status = "FAILED"
                source.failure_code = exc.code
                source.failure_message = str(exc)[:500]
            if job is not None:
                job.status = "FAILED"
                job.error_code = exc.code
                job.error_message = str(exc)[:500]
                job.completed_at = utcnow()
            session.commit()
        except Exception:
            session.rollback()
            source = session.get(SupportAIKnowledgeSourceRow, source_id)
            job = session.get(SupportAIIngestionJobRow, job_id)
            if source is not None:
                if source.approved_at is None or source.chunk_count <= 0:
                    source.status = "FAILED"
                source.failure_code = "KNOWLEDGE_INGESTION_FAILED"
                source.failure_message = "知识文件处理失败，请稍后重试。"
            if job is not None:
                job.status = "FAILED"
                job.error_code = "KNOWLEDGE_INGESTION_FAILED"
                job.error_message = "知识文件处理失败，请稍后重试。"
                job.completed_at = utcnow()
            session.commit()
